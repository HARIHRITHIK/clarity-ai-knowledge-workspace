"""
Generates multi-format sample dataset in sample_data/
Used for automated testing, benchmarks, and offline recruiter demonstrations.
"""

from pathlib import Path
import docx
from fpdf import FPDF

OUT_DIR = Path(__file__).resolve().parent.parent / "sample_data"
OUT_DIR.mkdir(exist_ok=True)

# 1. Generate sample_data/project_notes.txt
notes_content = """ACME CORPORATION - SPRINT & ARCHITECTURE NOTES
Date: February 14, 2025
Attendees: David Kim (CTO), Michael Chen (CEO), Sarah Rodriguez (CFO), Priya Mehta (Engineering Lead)

1. Executive Summary
The engineering team completed the migration to the v3.0 API infrastructure. Total response latency dropped by 42%, while memory overhead remained within budgeted thresholds.

2. Key Action Items
- David Kim to finalize the multi-region failover architecture by March 1, 2025.
- Priya Mehta to conduct load testing on the indexing pipeline with 50,000 concurrent documents.
- Sarah Rodriguez to approve the $180,000 infrastructure allocation for Q2 2025.

3. Vendor & Partner Updates
TechSupply Solutions delivered the updated service level agreement (SLA) with guaranteed 99.95% uptime.
"""
(OUT_DIR / "project_notes.txt").write_text(notes_content.strip(), encoding="utf-8")

# 2. Generate sample_data/sales_summary.csv
csv_content = """Date,Customer,Segment,Amount,Status,Sales Representative
2025-01-15,Meridian Financial Group,Enterprise SaaS,$84000,Completed,Robert Martinez
2025-01-18,Northstar Logistics,Enterprise SaaS,$68000,Completed,Amanda Foster
2025-01-22,BrightPath Healthcare,Enterprise SaaS,$92000,Completed,Robert Martinez
2025-02-01,Apex Digital,Mid-Market,$18500,Completed,Lisa Thompson
2025-02-05,Coastal Energy Partners,Enterprise SaaS,$115000,Completed,Amanda Foster
"""
(OUT_DIR / "sales_summary.csv").write_text(csv_content.strip(), encoding="utf-8")

# 3. Generate sample_data/employee_policy.docx
doc = docx.Document()
doc.add_heading("Acme Corporation - Global Employee Policy Handbook", 0)
doc.add_paragraph("Effective Date: January 1, 2025 | Approved by: People Operations & HR")

p1 = doc.add_paragraph(
    "Acme Corporation maintains a modern, flexible work policy designed to support high performance and work-life balance. "
    "Full-time employees receive 20 days of paid time off (PTO) annually, alongside standard public holidays and health benefits."
)

doc.add_heading("Remote & Hybrid Work Guidelines", level=1)
doc.add_paragraph(
    "Employees in eligible engineering and analysis roles may work remotely up to 3 days per week. "
    "Core collaboration hours are 10:00 AM to 4:00 PM in your local time zone."
)

doc.add_heading("Professional Development & Compensation", level=1)
doc.add_paragraph(
    "Each employee is entitled to an annual education stipend of $2,500 for approved courses, conferences, and technical certifications. "
    "Annual performance reviews occur in November, managed by Vice President of People Operations Eleanor Grant."
)
doc.save(str(OUT_DIR / "employee_policy.docx"))

# 4. Generate sample_data/company_report.pdf
pdf = FPDF(format="A4")
pdf.add_page()
pdf.set_font("Helvetica", "B", 16)
pdf.cell(190, 10, "Acme Corporation - Q4 Financial Report", align="C", new_x="LMARGIN", new_y="NEXT")
pdf.set_font("Helvetica", "", 10)
pdf.cell(190, 8, "Fiscal Year Ending December 31, 2024 | Prepared by Sarah Rodriguez, CFO", align="C", new_x="LMARGIN", new_y="NEXT")
pdf.ln(5)

pdf.set_font("Helvetica", "B", 12)
pdf.cell(190, 8, "Executive Financial Summary", new_x="LMARGIN", new_y="NEXT")
pdf.set_font("Helvetica", "", 10)
summary_text = (
    "Acme Corporation closed fiscal year 2024 with total annual revenue of $24.7 million, representing an 18% "
    "increase over fiscal year 2023. Fourth quarter revenue alone reached $7.1 million. Net profit margin improved "
    "to 14.3%, reflecting strong operational efficiency in the Enterprise SaaS segment led by CEO Michael Chen."
)
pdf.multi_cell(190, 6, summary_text)
pdf.ln(4)

pdf.set_font("Helvetica", "B", 12)
pdf.cell(190, 8, "Operational Highlights & Key Facts", new_x="LMARGIN", new_y="NEXT")
pdf.set_font("Helvetica", "", 10)
facts = [
    "- Enterprise SaaS annual revenue reached $16.2 million (65.6% of total revenue).",
    "- Customer Net Retention rate stood at 112%, with annual churn reduced to 5.4%.",
    "- Engineering investment was $4.2 million under CTO David Kim, shipping 4 major product releases.",
    "- Total cash reserves as of December 31, 2024 totaled $6.8 million with zero long-term debt.",
]
for fact in facts:
    pdf.multi_cell(190, 6, fact)

pdf.output(str(OUT_DIR / "company_report.pdf"))
print("Successfully generated all 4 files in sample_data/")
